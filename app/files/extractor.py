from __future__ import annotations

import io
import csv
from pathlib import Path
from typing import Any

from app.files.models import DocumentContent, DocumentSection, DocumentTable
from app.files.tokens import estimate_tokens, estimate_tokens_for_tables, categorize_size

# Supported extensions per AGENTS §4
SUPPORTED_TYPES = {".txt", ".md", ".csv", ".xlsx", ".docx", ".pdf"}

SUPPORTED_TYPES_STR = ", ".join(sorted(SUPPORTED_TYPES))


def detect_file_type(filename: str) -> str:
    """Return normalized file type (without dot) or raise for unsupported."""
    ext = Path(filename).suffix.lower()
    if ext not in SUPPORTED_TYPES:
        raise ValueError(f"Unsupported file type {ext!r}. Supported: {SUPPORTED_TYPES_STR}")
    # md -> md, txt -> txt etc; keep as extension without dot
    return ext.lstrip(".")


def _decode_text_bytes(data: bytes) -> str:
    # Try utf-8, fallback to latin1 with errors replace; never fail
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def extract_txt(data: bytes, filename: str) -> DocumentContent:
    text = _decode_text_bytes(data)
    # normalize line endings
    text = text.replace("\r\n", "\n").strip()
    doc = DocumentContent(filename=filename, file_type="txt", text=text, metadata={"bytes": len(data)})
    return doc


def extract_md(data: bytes, filename: str) -> DocumentContent:
    text = _decode_text_bytes(data).replace("\r\n", "\n").strip()
    doc = DocumentContent(filename=filename, file_type="md", text=text, metadata={"bytes": len(data)})
    return doc


def extract_csv(data: bytes, filename: str) -> DocumentContent:
    text_raw = _decode_text_bytes(data)
    # Use csv module to keep lightweight; pandas alternative not required for small CSV
    # But try pandas if available for better type handling
    headers: list[str] = []
    rows: list[list[Any]] = []
    text_parts: list[str] = []
    try:
        # Use pandas for robust parsing if available
        import pandas as pd  # type: ignore

        df = pd.read_csv(io.BytesIO(data), dtype=str, keep_default_na=False, engine="python")
        # pandas may infer; keep as strings
        headers = [str(c).strip() for c in df.columns.tolist()]
        # Replace NaN with ""
        df = df.fillna("")
        rows = df.astype(str).values.tolist()
        # Build text representation
        text_parts.append(" | ".join(headers))
        for r in rows:
            text_parts.append(" | ".join(str(c) for c in r))
        text = "\n".join(text_parts)
    except Exception:
        # fallback to csv
        reader = csv.reader(io.StringIO(text_raw))
        all_rows = list(reader)
        if not all_rows:
            text = ""
        else:
            headers = [h.strip() for h in all_rows[0]]
            rows = all_rows[1:]
            text_parts = []
            if headers:
                text_parts.append(" | ".join(headers))
            for r in rows:
                text_parts.append(" | ".join(str(c) for c in r))
            text = "\n".join(text_parts)

    # Truncate very long cells per §29
    truncated = False
    max_cell = 4000
    for i, h in enumerate(headers):
        if len(h) > max_cell:
            headers[i] = h[:max_cell] + "…"
            truncated = True
    for r in rows:
        for j, c in enumerate(r):
            s = str(c)
            if len(s) > max_cell:
                r[j] = s[:max_cell] + "…"
                truncated = True

    table = DocumentTable(
        name=Path(filename).stem,
        headers=headers,
        rows=rows,
        source=f"{filename} Rows: 1-{len(rows)}" if rows else filename,
        metadata={"bytes": len(data), "truncated": truncated},
    )
    # For CSV, text is the flattened table; also store tables list
    doc = DocumentContent(
        filename=filename,
        file_type="csv",
        text=text,
        tables=[table] if headers or rows else [],
        metadata={"bytes": len(data), "columns": headers},
    )
    return doc


def extract_xlsx(data: bytes, filename: str) -> DocumentContent:
    try:
        import openpyxl  # type: ignore
    except ImportError as e:
        raise RuntimeError("openpyxl is required for .xlsx extraction") from e

    wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    sheets: list[str] = []
    tables: list[DocumentTable] = []
    texts: list[str] = []
    # Hard caps per finding: protect 12GB laptops from huge XLSX
    MAX_TOTAL_ROWS = 5000
    MAX_TOTAL_CHARS = 200_000  # ~50k tokens
    total_rows = 0
    total_chars = 0
    truncated = False
    trunc_reason: str | None = None

    def _row_chars(vals: list[Any]) -> int:
        return sum(len(str(v)) if v is not None else 0 for v in vals) + len(vals)  # separators

    outer_truncated = False
    for ws in wb.worksheets:
        if total_rows >= MAX_TOTAL_ROWS or total_chars >= MAX_TOTAL_CHARS:
            truncated = True
            if not trunc_reason:
                trunc_reason = "rows" if total_rows >= MAX_TOTAL_ROWS else "chars"
            break
        sheets.append(ws.title)
        rows_iter = ws.iter_rows(values_only=True)
        all_rows: list[list[Any]] = []
        for row in rows_iter:
            if total_rows >= MAX_TOTAL_ROWS or total_chars >= MAX_TOTAL_CHARS:
                truncated = True
                if not trunc_reason:
                    trunc_reason = "rows" if total_rows >= MAX_TOTAL_ROWS else "chars"
                outer_truncated = True
                break
            # row is tuple of values; normalize None -> ""
            # skip entirely empty rows at end? Keep but skip if all None/empty
            if all(v is None or (isinstance(v, str) and not v.strip()) for v in row):
                continue
            # trim trailing None/empty
            # convert to list of strings/next
            vals = [("" if v is None else v) for v in row]
            # remove trailing empties
            while vals and (vals[-1] == "" or vals[-1] is None):
                vals.pop()
            if not vals:
                continue
            # Cap checks before appending
            rc = _row_chars(vals)
            if total_rows + 1 > MAX_TOTAL_ROWS or total_chars + rc > MAX_TOTAL_CHARS:
                truncated = True
                if not trunc_reason:
                    trunc_reason = "rows" if total_rows + 1 > MAX_TOTAL_ROWS else "chars"
                outer_truncated = True
                break
            all_rows.append(vals)
            total_rows += 1
            total_chars += rc
        if outer_truncated:
            # still need to break outer sheets loop after handling current sheet's partial data
            # if we broke mid-sheet, we have partial all_rows to process as table below
            pass
        else:
            # check chars for sheet header text overhead (approx)
            total_chars += len(ws.title) + 10

        if not all_rows:
            continue

        # Assume first row is headers if it contains strings and second row differs type? Keep simple: first row as headers
        headers = [str(c).strip() if c is not None else "" for c in all_rows[0]]
        data_rows = all_rows[1:] if len(all_rows) > 1 else []

        # If headers look like data (e.g., numeric) and no clear header, we still treat first row as headers for structure
        # Provide option to keep headers even if empty?

        # Truncate long cells
        max_cell = 4000
        headers = [h[:max_cell] + "…" if len(h) > max_cell else h for h in headers]
        for r in data_rows:
            for j, c in enumerate(r):
                s = str(c) if c is not None else ""
                if len(s) > max_cell:
                    r[j] = s[:max_cell] + "…"

        table = DocumentTable(
            name=ws.title,
            headers=headers,
            rows=[[str(c) if c is not None else "" for c in r] for r in data_rows],
            source=f"{filename} Sheet: {ws.title} Rows: 1-{len(data_rows)}" if data_rows else f"{filename} Sheet: {ws.title}",
            metadata={"sheet": ws.title},
        )
        tables.append(table)
        # Build text for this sheet
        texts.append(f"Sheet: {ws.title}")
        if headers:
            texts.append(" | ".join(headers))
        for r in data_rows:
            texts.append(" | ".join(str(c) if c is not None else "" for c in r))
        texts.append("")  # blank between sheets
        if outer_truncated:
            break

    text = "\n".join(texts).strip()
    # Ensure capped text also respects char cap (safety)
    if len(text) > MAX_TOTAL_CHARS:
        text = text[:MAX_TOTAL_CHARS] + "…[truncated]"
        truncated = True
        if not trunc_reason:
            trunc_reason = "chars"
    doc = DocumentContent(
        filename=filename,
        file_type="xlsx",
        text=text,
        tables=tables,
        sheets=sheets,
        metadata={
            "bytes": len(data),
            "truncated": truncated,
            "truncation_reason": trunc_reason,
            "total_rows": total_rows,
            "total_chars": total_chars,
        },
    )
    try:
        wb.close()
    except Exception:
        pass
    return doc


def extract_docx(data: bytes, filename: str) -> DocumentContent:
    try:
        import docx  # type: ignore

        docx_bytes = io.BytesIO(data)
        document = docx.Document(docx_bytes)
    except ImportError as e:
        raise RuntimeError("python-docx is required for .docx extraction") from e
    except Exception as e:
        raise ValueError(f"Failed to parse docx {filename!r}: {e}") from e

    sections: list[DocumentSection] = []
    full_text_parts: list[str] = []
    current_title = "Document"
    current_content: list[str] = []
    # Heading detection via style names
    def is_heading(paragraph) -> tuple[bool, str]:
        style_name = paragraph.style.name if paragraph.style else ""
        # docx styles like Heading 1, Heading 2, Title
        if style_name.startswith("Heading"):
            return True, paragraph.text.strip()
        if style_name == "Title":
            return True, paragraph.text.strip()
        # Also heuristic: all caps short line? skip
        return False, ""

    for para in document.paragraphs:
        txt = para.text.strip()
        if not txt:
            # blank paragraph separates but keep as newline
            if current_content:
                current_content.append("")
            continue
        heading, title = is_heading(para)
        if heading and title:
            # flush previous section
            if current_content or current_title != "Document":
                content = "\n".join(current_content).strip()
                if content or current_title != "Document":
                    sections.append(
                        DocumentSection(
                            title=current_title,
                            content=content,
                            source=f"{filename} Section: {current_title}",
                        )
                    )
                    if content:
                        full_text_parts.append(f"## {current_title}\n{content}")
                current_content = []
            current_title = title
        else:
            current_content.append(txt)

    # flush last
    if current_content or current_title:
        content = "\n".join(current_content).strip()
        # avoid duplicate empty first section if no headings
        if content or not sections:
            # if no sections yet, this is the whole doc under Document
            if not sections and current_title == "Document":
                sections.append(
                    DocumentSection(title=current_title, content=content, source=f"{filename} Section: {current_title}")
                )
            else:
                # only add if content non-empty or title changed
                if content or current_title != "Document":
                    sections.append(
                        DocumentSection(
                            title=current_title,
                            content=content,
                            source=f"{filename} Section: {current_title}",
                        )
                    )
            if content:
                if len(sections) == 1 and sections[0].title == "Document":
                    full_text_parts.append(content)
                else:
                    full_text_parts.append(f"## {current_title}\n{content}")

    # Tables in docx
    tables: list[DocumentTable] = []
    for idx, tbl in enumerate(document.tables):
        headers: list[str] = []
        rows: list[list[str]] = []
        tbl_rows = list(tbl.rows)
        if not tbl_rows:
            continue
        # assume first row headers
        header_cells = tbl_rows[0].cells
        headers = [c.text.strip() for c in header_cells]
        for r in tbl_rows[1:]:
            rows.append([c.text.strip() for c in r.cells])
        table = DocumentTable(
            name=f"Table {idx+1}",
            headers=headers,
            rows=rows,
            source=f"{filename} Table: {idx+1}",
            metadata={"table_index": idx},
        )
        tables.append(table)
        # add to text
        full_text_parts.append(f"Table {idx+1}: " + " | ".join(headers))
        for r in rows:
            full_text_parts.append(" | ".join(r))

    text = "\n\n".join(p for p in full_text_parts if p).strip()
    # If no sections/tables, fallback to plain paragraphs join
    if not text:
        # raw fallback: join all paragraph texts
        text = "\n".join(p.text for p in document.paragraphs if p.text.strip())

    doc = DocumentContent(
        filename=filename,
        file_type="docx",
        text=text,
        tables=tables,
        sections=sections,
        metadata={"bytes": len(data), "paragraphs": len(document.paragraphs)},
    )
    return doc


def extract_pdf(data: bytes, filename: str) -> DocumentContent:
    try:
        try:
            import pymupdf as fitz  # type: ignore  # PyMuPDF >=1.24
        except ImportError:
            import fitz  # type: ignore  # fallback

        doc = fitz.open(stream=data, filetype="pdf")
    except ImportError as e:
        raise RuntimeError("PyMuPDF is required for .pdf extraction") from e
    except Exception as e:
        raise ValueError(f"Failed to parse pdf {filename!r}: {e}") from e

    pages_text: list[str] = []
    full_parts: list[str] = []
    for i, page in enumerate(doc):
        txt = page.get_text("text") or ""
        txt = txt.strip()
        pages_text.append(txt)
        if txt:
            # annotate with page citation
            full_parts.append(f"[Page {i+1}]\n{txt}")
    text = "\n\n".join(full_parts).strip()
    # If no text extracted, try alternative (maybe scanned)
    if not text:
        text = ""

    result = DocumentContent(
        filename=filename,
        file_type="pdf",
        text=text,
        pages=len(doc),
        metadata={"bytes": len(data), "pages": len(doc)},
    )
    # Optionally store per-page as sections for citation
    # We'll create sections per page if needed for chunking
    sections: list[DocumentSection] = []
    for idx, pt in enumerate(pages_text):
        if pt.strip():
            sections.append(
                DocumentSection(
                    title=f"Page {idx+1}",
                    content=pt.strip(),
                    source=f"{filename} Page: {idx+1}",
                )
            )
    result.sections = sections
    try:
        doc.close()
    except Exception:
        pass
    return result


def extract_document(file_bytes: bytes, filename: str) -> DocumentContent:
    """Main entry: detect type, dispatch extraction, then normalize sizing.

    Follows AGENTS §3 pipeline: detection -> extraction -> normalized -> size check.
    """
    file_type = detect_file_type(filename)
    # Dispatch
    if file_type in ("txt",):
        doc = extract_txt(file_bytes, filename)
    elif file_type == "md":
        doc = extract_md(file_bytes, filename)
    elif file_type == "csv":
        doc = extract_csv(file_bytes, filename)
    elif file_type == "xlsx":
        doc = extract_xlsx(file_bytes, filename)
    elif file_type == "docx":
        doc = extract_docx(file_bytes, filename)
    elif file_type == "pdf":
        doc = extract_pdf(file_bytes, filename)
    else:
        # Should not reach due to detect
        raise ValueError(f"Unsupported file type {file_type!r}")

    # Size check per §8: use extracted text length / estimated tokens
    # Combine text + tables for token estimate
    base_tokens = estimate_tokens(doc.text)
    table_tokens = estimate_tokens_for_tables(doc.tables)
    # sections already counted in text but include if docx/pdf sections extra?
    total_tokens = base_tokens + table_tokens
    # For docx/pdf where sections duplicate text, avoid double-count: if doc.text already includes sections, don't add double
    # Heuristic: if doc.text length > sections total, use base+table only
    # Otherwise keep total as base (which is text) — we already did base+table, which may double if tables duplicated in text for xlsx/csv
    # For xlsx/csv, text is derived from tables, so base_tokens ~ table_tokens; we should take max not sum for those
    if file_type in ("xlsx", "csv"):
        # text derived from tables, use max to avoid double
        total_tokens = max(base_tokens, table_tokens) if table_tokens else base_tokens
    else:
        # For docx/pdf, sections are included in text, so not adding extra
        # keep as base + table (table may be extra)
        pass

    doc.token_estimate = total_tokens
    doc.size_category = categorize_size(total_tokens)
    # Enrich metadata
    doc.metadata["token_estimate"] = total_tokens
    doc.metadata["size_category"] = doc.size_category
    return doc
